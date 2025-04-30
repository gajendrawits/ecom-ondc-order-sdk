from docx import Document
from docx.shared import Pt

doc = Document()

def add_heading(text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    run = doc.add_paragraph().add_run(text)
    run.bold = bold
    run.font.size = Pt(11)

# Title
add_heading("RESUME", level=0)

# Personal Information
add_heading("Personal Information", level=1)
add_paragraph("Name: Dinesh Singh Ramola")
add_paragraph("Father's Name: Trilok Singh Ramola")
add_paragraph("Date of Birth: 23 March 1979")
add_paragraph("Marital Status: Married")
add_paragraph("Height: 182 cm")
add_paragraph("Blood Group: A+ve")
add_paragraph("Birth Place: Chamoli, Uttarakhand")
add_paragraph("Caste: Rajput")
add_paragraph("Religion: Hindu")
add_paragraph("Nationality: Indian")
add_paragraph("Mobile No: 8006400150, 8006300707")
add_paragraph("Email: deenuramola7923@gmail.com")
add_paragraph("Address: Shivlok Enclave Telpur Chowk, PO: Mehuwala Mafi, Teh: Dehra Dun, Dist: Dehra Dun, State: Uttarakhand, Pin: 248018")

# Career Summary
add_heading("Professional Summary", level=1)
add_paragraph("Served in the Indian Army for 28 years as a Radio Telephone Operator, with expertise in handling communication equipment, weapons, and providing high-level security.")

# Skills
add_heading("Skills", level=1)
skills = [
    "Handling all types of Army communication equipment",
    "Radio sets, Motorola sets, Army Exchanges, Telephone sets",
    "Telephone line laying and repair",
    "Handling all security equipment and weapons",
    "Special security duties for Brigade Commander (3 years in Jammu & Kashmir)"
]
for skill in skills:
    add_paragraph(f"- {skill}")

# Work Experience
add_heading("Work Experience", level=1)
add_paragraph("Trade: Radio Telephone Operator")
add_paragraph("Date of Enrolment: 25 February 1997")
add_paragraph("Date of Retirement: 28 February 2025")
add_paragraph("Postings: Arty Centre Nashik Road, Assam, Jammu & Kashmir (multiple), Jodhpur, Sagar, Dehra Dun, Ajmer, Faridkot")

# Education
add_heading("Education", level=1)
education = [
    "10th - GHS Harmani Chamoli, Uttarakhand - First Division - 1996",
    "12th - Janta Inter College Naya Gaon, Dehra Dun - Second Division - 2004",
    "Graduate - DAV (PG) College, Dehra Dun - Qualified 2005",
    "2nd Year - SGRR (PG) College, Dehra Dun - Qualified 2006",
    "Final Year - CNI Degree College, Dehra Dun - Qualified 2007"
]
for edu in education:
    add_paragraph(f"- {edu}")

# Languages
add_heading("Languages", level=1)
add_paragraph("Hindi, Garhwali")

# Awards and Achievements
add_heading("Awards & Achievements", level=1)
awards = [
    "Received multiple medals in sports",
    "Certificate of Best Performing Combat Assault Team from 15 Corps Battle School, Jammu & Kashmir",
    "Appreciation Certificate from Commander for Best Team Leader in Security Duties in Jammu & Kashmir"
]
for award in awards:
    add_paragraph(f"- {award}")

# Hobbies
add_heading("Hobbies", level=1)
add_paragraph("Playing music, Basketball, Volleyball, Handball, Hockey")

# Weapon License
add_heading("Weapon License", level=1)
add_paragraph("Weapon: Pistol 7.65 mm")
add_paragraph("License: All India")
add_paragraph("License Issued From: Dehra Dun, Uttarakhand")

# Save the file
doc.save("Dinesh_Singh_Ramola_Resume.docx")
